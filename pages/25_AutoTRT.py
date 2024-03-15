import numpy as np
import pandas as pd
import streamlit as st
import datetime as datetime
from datetime import datetime as dt
import plotly.express as px
import csv
from scipy.stats import linregress
from docx import Document
from docx.shared import Cm
from docx.enum.style import WD_STYLE_TYPE
import io

st.set_page_config(page_title="TRT-beregning", page_icon="🔥")

with open("styles/main.css") as f:
    st.markdown("<style>{}</style>".format(f.read()), unsafe_allow_html=True)

class TRT_beregning:
    def __init__(self):
        pass
    
    def kjor_hele(self):
        self.streamlit_input()
        if self.datafil and (self.knapp == True):
            self.varmeegenskaper()
            self.les_av_datafil()
            self.finn_denne_test()
            self.del_test()
            self.linear_tiln()
            self.plot1()
            self.effektiv_varmeledningsevne()
            self.plot2()
            self.teoretiske_tempforlop()
            self.plot3()
            self.plot4()
            self.plot5()
            self.tilfort_effekt()
            self.plot6()
            self.lagre_som_png()
            self.lag_rapport()

    def streamlit_input(self):
        # Viser alle input-felt i streamlit
        st.title('Resultatberegning fra termisk responstest (TRT)')
        st.header('Inndata')

        self.datafil = st.file_uploader(label='Datafil fra testrigg (.dat)',type='dat')

        c1, c2 = st.columns(2)
        with c1:
            self.kollvaeske = st.selectbox(label='Type kollektorvæske',options=['HXi24','HXi35','Kilfrost Geo 24 %','Kilfrost Geo 32 %','Kilfrost Geo 35 %'],index=3)
        #with c2:
        #    pass

        c1, c2 = st.columns(2)
        with c1:
            self.dybde = st.number_input(label='Brønndybde (m)', value=250, step=1)
        with c2:
            self.diam = st.number_input(label='Diameter (mm)', value=115, step=1)

        # Streamlit-input av datoer:
        c1, c2 = st.columns(2)
        with c1:
            self.st_startdato = st.date_input(label='Startdato for termisk responstest', value=dt.now(), max_value=dt.now())
        with c2:
            self.st_sluttdato = st.date_input(label='Sluttdato for termisk responstest', value=dt.now(), max_value=dt.now())

        c1, c2 = st.columns(2)
        with c1:
            self.strom_foer = st.number_input(label='Strømmåler før test (kWh)', value=20000, step=1)
        with c2:
            self.strom_etter = st.number_input(label='Strømmåler etter test (kWh)', value=21000, step=1)

        st.caption('')
        self.knapp = st.button('Kjør')

    def varmeegenskaper(self):
        if self.kollvaeske == 'HXi24':
            self.tetthet = 970.5
            self.varmekap = 4.298
        elif self.kollvaeske == 'HXi35':
            self.tetthet = 955
            self.varmekap = 4.061
        elif self.kollvaeske == 'Kilfrost Geo 24 %':
            self.tetthet = 1105.5
            self.varmekap = 3.455 #kJ/kg K
        elif self.kollvaeske == 'Kilfrost Geo 32 %':
            self.tetthet = 1136.2
            self.varmekap = 3.251 #kJ/kg K
        elif self.kollvaeske == 'Kilfrost Geo 35 %':
            self.tetthet = 1150.6
            self.varmekap = 3.156 #kJ/kg K

    def les_av_datafil(self):
        
        @st.cache_data
        def funk_les_datafil(filnavn):
            df = pd.read_csv(filnavn, 
                    sep=",", 
                    skiprows=1, 
                    usecols=[0,6,7,8,9,10,11,12,13,23],
                    header=0,
                    names = ['tidspunkt','temp_fra_bronn','aPump','temp_til_bronn','rigg','ute','sirk_hast','pumpehast','panelhast','avpaa'])
            
            df = df.drop([0,1])
            df['temp_fra_bronn'] = df['temp_fra_bronn'].astype(float)
            df['aPump'] = df['aPump'].astype(float)
            df['temp_til_bronn'] = df['temp_til_bronn'].astype(float)
            df['rigg'] = df['rigg'].astype(float)
            df['ute'] = df['ute'].astype(float)
            df['sirk_hast'] = df['sirk_hast'].astype(float)
            df['pumpehast'] = df['pumpehast'].astype(float)
            df['panelhast'] = df['panelhast'].astype(float)
            df['avpaa'] = df['avpaa'].astype(int)
            return df
        
        self.df = funk_les_datafil(self.datafil)

    def finn_denne_test(self):
        # Gjør datoer om til datetime-format;
        def custom_to_datetime(timestamp):
            try:
                return pd.to_datetime(timestamp, format="%Y-%m-%d %H:%M:%S.%f")
            except ValueError:
                return pd.to_datetime(timestamp, format="%Y-%m-%d %H:%M:%S")

        self.df['tidspunkt'] = self.df['tidspunkt'].apply(custom_to_datetime)

        st_starttid = datetime.time(0, 0, 0)
        st_sluttid = datetime.time(23, 59, 59)

        startdato = datetime.datetime.combine(self.st_startdato,st_starttid)
        sluttdato = datetime.datetime.combine(self.st_sluttdato,st_sluttid)

        # Henter ut rader for kun de aktuelle datoer:
        denne_test = self.df.loc[(self.df['tidspunkt'] >= startdato) & (self.df['tidspunkt'] <= sluttdato)]
        denne_test = denne_test.reset_index(drop=True)

        # Henter ut rader kun for de aktuelle tidspunkter basert på hvor "Pump enabeled"-kolonnen slår inn:
        for i in range(0,len(denne_test)):
            if denne_test['avpaa'].iloc[i] != 0:
                startrad = int(i)
                break
        denne_test = denne_test.iloc[startrad:]
        denne_test = denne_test.reset_index(drop=True)

        # Regner ut gjennomsnittstemperaturen til kollektorvæsken, og legger dette som en egen kolonne:
        snittemp = pd.DataFrame((denne_test['temp_fra_bronn']+denne_test['temp_til_bronn'])/2)
        denne_test.insert(4,"snittemp",snittemp,True)

        self.denne_test = denne_test

    def del_test(self):

        # Deler inn i før og etter at varmen er slått på, basert på der det dukker opp to like klokkeslett:
        for i in range(0,len(self.denne_test)):
            if self.denne_test.iloc[i,0] == self.denne_test.iloc[i-1,0]:
                deleindeks = int(i)
                break

        test_del1 = self.denne_test.iloc[0:deleindeks,:]
        test_del2 = self.denne_test.iloc[deleindeks:,:]

        test_del1 = test_del1.reset_index(drop=True)                # Resetter radnummer til Dataframes
        test_del2 = test_del2.reset_index(drop=True)

        # Lager kolonne som viser antall sekunder siden testen startet
        startindeks = 0            #Raden hvor testen faktisk starter
        antall_sek = pd.DataFrame({'sek_siden_start' : [0]*len(test_del2)})
        for i in range(0,len(test_del2)):
            antall_sek.iloc[i] = (test_del2['tidspunkt'].iloc[i]-test_del2['tidspunkt'].iloc[startindeks]).total_seconds()

        test_del2.insert(1,"sek_siden_start",antall_sek,True)                  # Setter inn antall_sek som kolonne (indeks 1) i test_del2-dataframen
        test_del2['sek_siden_start'] = test_del2['sek_siden_start'].astype(float)

        # Definerer ln(t)
        ln_t = pd.DataFrame({'ln_t' : [0]*len(test_del2)})
        for i in range(0,len(test_del2)):
            ln_t.iloc[i] = np.log(test_del2['sek_siden_start'].iloc[i])
        
        test_del2.insert(2,"ln_t",ln_t,True)
        test_del2['ln_t'] = test_del2['ln_t'].astype(float)

        # Henter ut den delen av test_del2 som foregår etter 5 og 20 timer:
        etter5timer = test_del2.iloc[600:,:]                  #Antar her at det er et 30 sek mellom hvert målepunkt
        etter5timer = etter5timer.reset_index(drop=True)
        etter20timer = test_del2.iloc[2400:,:]                  #Antar her at det er et 30 sek mellom hvert målepunkt
        etter20timer = etter20timer.reset_index(drop=True)

        self.test_del1 = test_del1
        self.test_del2 = test_del2
        self.etter5timer = etter5timer
        self.etter20timer = etter20timer

    def linear_tiln(self):
        # Lineær tilnærming av gjennomsnittstemperaturen etter 20 timer:
        slope, intercept, r_value, p_value, std_err = linregress(self.etter20timer['ln_t'], self.etter20timer['snittemp'])
        self.y_pred = slope * self.etter20timer['ln_t'] + intercept
        self.r_verdi = str(round(r_value,3))

    def plot1(self):
        st.header('Resultater')

        til_plot1 = pd.DataFrame({"Tider" : self.etter20timer['ln_t'], "Gj.snittstemp. kollektorvæske" : self.etter20timer['snittemp'], "Lineær tilnærming med r = "+self.r_verdi+"" : self.y_pred})
        fig1 = px.line(til_plot1, x='Tider', y=["Gj.snittstemp. kollektorvæske","Lineær tilnærming med r = "+self.r_verdi+""], title='Temperaturmålinger etter 20 timer', color_discrete_sequence=['#367A2F', '#FFC358'])
        fig1.update_layout(xaxis_title='Logaritmen av tid siden teststart', yaxis_title='Temperatur (\u2103)',legend_title=None)
        fig1.update_xaxes(range=[11.0, 12.6], row=1, col=1)
        st.plotly_chart(fig1)
        self.fig1 = fig1
 
    def effektiv_varmeledningsevne(self):
        # Utregning av effektiv varmeledningsevne:
        self.indeks5timer = 600          # Dersom det er 30 sek mellom hver måling
        hjelpefunk = np.zeros(len(self.test_del2))
   
        for i in range(self.indeks5timer+1,len(self.test_del2)):
            y_values = np.array(self.test_del2['snittemp'].iloc[self.indeks5timer:i])
            x_values = np.array(self.test_del2['ln_t'].iloc[self.indeks5timer:i])
            hjelpefunk[i], intercept = np.polyfit(x_values, y_values, 1)
        
        tot_varighet_timer = self.test_del2['sek_siden_start'].iloc[-1]/3600
        mid_effekt = (self.strom_etter-self.strom_foer)/(tot_varighet_timer)*1000    # W
        
        effekt_per_m = (mid_effekt/self.dybde) # W
        ledn_evne = effekt_per_m/(4*np.pi*hjelpefunk)
        for i in range(0,len(ledn_evne)):                   # Fjerner urimelig høye verdier
            if ledn_evne[i] >= 20:
                ledn_evne[i] = float('nan')
        ledn_evne = pd.DataFrame(ledn_evne)

        self.stabil_ledn_evne = float(ledn_evne.tail(self.indeks5timer).mean())      # Antar at ledningsevnen stabiliserer seg på gjennomsnittet av de siste 5 timer

        self.mid_effekt = mid_effekt
        self.ledn_evne = ledn_evne

    def plot2(self):
        st.markdown("---")
        #self.ledn_evne_slider = st.slider('Varmeledningsevne som kurven konvergerer mot', 0.1, float(10), self.stabil_ledn_evne, step=0.01)
        self.ledn_evne_slider = st.number_input('Varmeledningsevne som kurven konvergerer mot. Velg verdien som passer best til plottet under.', 0.1, float(10), self.stabil_ledn_evne, step=0.01)
        stabil_ledn_evne_tilplot = pd.DataFrame({'Value' : [self.ledn_evne_slider]*len(self.test_del2)})
        til_plot2 = pd.DataFrame({"Tider" : self.test_del2['sek_siden_start'].iloc[self.indeks5timer:]/3600, "Effektiv ledningsevne" : self.ledn_evne[0].iloc[self.indeks5timer:], 'Stabilisert ledningsevne' : stabil_ledn_evne_tilplot['Value']})
        fig2 = px.line(til_plot2, x='Tider', y=['Effektiv ledningsevne','Stabilisert ledningsevne'], title='Utvikling av effektiv varmeledningsevne', color_discrete_sequence=['#367A2F', '#FFC358'])
        fig2.update_layout(xaxis_title='Tid siden teststart (timer)', yaxis_title='Effektiv varmeledningsevne (W/mK)',legend_title=None)
        fig2.update_xaxes(range=[0, 75], row=1, col=1)
        fig2.update_yaxes(range=[2, 5], row=1, col=1)
        st.plotly_chart(fig2)
        self.fig2 = fig2
        
    def teoretiske_tempforlop(self):    
        # Regner ut teoretiske temperaturforløp med gitte verdier av borehullsmotstand:
        snitt_temp_inn = self.test_del1.loc[10:, 'temp_til_bronn'].mean()
        snitt_temp_ut = self.test_del1.loc[10:, 'temp_fra_bronn'].mean()                      #Bruker alle temp. målinger unntatt de 10 første
        self.uforst_temp_verdi = float((snitt_temp_inn+snitt_temp_ut)/2)

        lambdaa = self.ledn_evne_slider             # Den verdien som kurven i fig 2 stabiliserer seg mot.
        varmekap_Jm3K = 2200000
        diff = lambdaa/varmekap_Jm3K 
        #self.motstand = st.slider('Varmemotstand (mK/W)', float(0), float(1), 0.08, step=0.01)
        I1 = (self.mid_effekt/(4*np.pi*lambdaa*self.dybde))
        motstand_gjett_vektor = (np.array(self.test_del2['snittemp'])-self.uforst_temp_verdi-I1*np.array(self.test_del2['ln_t'])-I1*(np.log((4*diff)/(((self.diam/1000)/2)**2))-0.5772)) * (self.dybde/self.mid_effekt)
        motstand_gjett = np.mean(motstand_gjett_vektor[3:])

        st.markdown("---")
        self.motstand = st.number_input('Varmemotstand (mK/W). Velg den verdien som gir best samsvar med kurven under.', float(0), float(1), motstand_gjett, step=0.001)

        teori_temp = self.uforst_temp_verdi+I1*np.array(self.test_del2['ln_t'])+I1*(np.log((4*diff)/(((self.diam/1000)/2)**2))-0.5772)+(self.mid_effekt/self.dybde)*motstand_gjett
        self.teori_temp = pd.DataFrame(teori_temp)

    def plot3(self):
        motstand_str = str(round(self.motstand,3))
        snittemp_tilplot = pd.DataFrame({"Tider" : self.test_del2['sek_siden_start']/3600, "Gj.snittstemp. kollektorvæske" : self.test_del2['snittemp'], "Typekurve R = "+motstand_str+" mK/W" : self.teori_temp[0]})
        fig3 = px.line(snittemp_tilplot, x='Tider', y=["Gj.snittstemp. kollektorvæske","Typekurve R = "+motstand_str+" mK/W"], title='Faktisk temp.forløp og typekurve for termisk motstand R = '+motstand_str+' mK/W', color_discrete_sequence=['#367A2F', '#FFC358'])
        fig3.update_layout(xaxis_title='Tid siden teststart (timer)', yaxis_title='Temperatur (\u2103)',legend_title=None)
        st.plotly_chart(fig3)
        self.fig3 = fig3

    def plot4(self):
        st.markdown("---")
        # Definerer uforstyrret temperatur:
        uforst_temp = pd.DataFrame({'Value' : [self.uforst_temp_verdi]*len(self.test_del1)})
        
        df_tilplot = pd.DataFrame({"Tider" : self.test_del1['tidspunkt'], "Gj.snittstemp. kollektorvæske" : self.test_del1['temp_fra_bronn'], "Uforstyrret temperatur" : uforst_temp['Value']})
        fig4 = px.line(df_tilplot, x='Tider', y=["Gj.snittstemp. kollektorvæske","Uforstyrret temperatur"], title='Temperaturmålinger fra sirkulasjon av kollektorvæske', color_discrete_sequence=['#367A2F', '#FFC358'])
        fig4.update_layout(xaxis_title='Tid (klokkeslett)', yaxis_title='Temperatur (\u2103)',legend_title=None)
        st.plotly_chart(fig4)
        self.fig4 = fig4

    def plot5(self):
        st.markdown("---")
        til_plot5 = pd.DataFrame({"Tider" : self.test_del2['sek_siden_start']/3600, "Temperatur til brønnen" : self.test_del2['temp_til_bronn'], "Temperatur fra brønnen" : self.test_del2['temp_fra_bronn'], "Temperatur i testrigg" : self.test_del2['rigg'], "Temperatur i uteluft" : self.test_del2['ute']})
        fig5 = px.line(til_plot5, x='Tider', y=['Temperatur til brønnen', 'Temperatur fra brønnen', 'Temperatur i testrigg', 'Temperatur i uteluft'], title='Utelufttemp. og kollektorvæsketemp. til og fra energibrønnen', color_discrete_sequence=['#367A2F', '#C2CF9F', '#FFC358', '#FFE7BC'])
        fig5.update_layout(xaxis_title='Tid fra teststart (timer)', yaxis_title='Temperatur (\u2103)',legend_title=None)
        st.plotly_chart(fig5)
        self.fig5 = fig5

    def tilfort_effekt(self):
        # Regner ut tilført varmeeffekt ved bruk av sirkulasjonshastighet:
        tilfort_eff = self.tetthet*(np.array(self.test_del2['sirk_hast'])/60)/1000*self.varmekap*np.abs(np.array(self.test_del2['temp_til_bronn'])-np.array(self.test_del2['temp_fra_bronn']))
        self.tilfort_eff = pd.DataFrame(tilfort_eff)

    def plot6(self):
        st.markdown("---")
        til_plot6 = pd.DataFrame({"Tider" : self.test_del2['sek_siden_start']/3600, "Tilført varmeeffekt" : self.tilfort_eff[0], "Sirkulasjonshastighet" : self.test_del2['sirk_hast']})
        fig6 = px.scatter(til_plot6, x='Tider', y=["Tilført varmeeffekt","Sirkulasjonshastighet"], title='Sirkulasjonshastighet og tilført varmeeffekt', color_discrete_sequence=['#367A2F', '#FFC358'])
        fig6.update_layout(xaxis_title='Tid siden teststart (timer)', yaxis_title='Sirkulasjonsmengde [l/min] og tilført effekt [kW]',legend_title=None)
        st.plotly_chart(fig6)
        self.fig6 = fig6

    def lagre_som_png(self):
        fig_bredde = 800
        fig_hoyde = 450
        self.fig1.write_image("src/data/TRT-figurer/fig1.png",width=fig_bredde, height=fig_hoyde)
        self.fig2.write_image("src/data/TRT-figurer/fig2.png",width=fig_bredde, height=fig_hoyde)
        self.fig3.write_image("src/data/TRT-figurer/fig3.png",width=fig_bredde, height=fig_hoyde)
        self.fig4.write_image("src/data/TRT-figurer/fig4.png",width=fig_bredde, height=fig_hoyde)
        self.fig5.write_image("src/data/TRT-figurer/fig5.png",width=fig_bredde, height=fig_hoyde)
        self.fig6.write_image("src/data/TRT-figurer/fig6.png",width=fig_bredde, height=fig_hoyde)

    def lag_rapport(self):
        sted = "Hos meg"
        oppdragsgiver = "Askeladden"
        document = Document("src/data/TRT-notat/Mal Rapport TRT - .docx")
        styles = document.styles
        style = styles.add_style('Citation', WD_STYLE_TYPE.PARAGRAPH)
        #document.paragraphs[1].text = f"Oppdragsgiver - {oppdragsgiver}"
        #document.paragraphs[2].text = f"Tittel på rapport: Termisk responstest - {sted}"
        #document.paragraphs[3].text = f"Oppdragsnavn: Geoteknisk rapport - {sted}"
        #document.paragraphs[4].text = f"Oppdragsnummer: - 123456789""
        #document.add_heading("Innledning", 1)
        #document.add_paragraph(report_text_1)
        
        for paragraph_index, paragraph in enumerate(document.paragraphs):
            if "[python_fig2]" in paragraph.text:
                linje_til_fig2 = paragraph_index
                break  
        document.paragraphs[linje_til_fig2].clear()
        run2 = document.paragraphs[linje_til_fig2].add_run()
        bilde2 = run2.add_picture("src/data/TRT-figurer//fig2.png")
        bilde2.width = Cm(16)
        bilde2.height = Cm(9)

        for paragraph_index, paragraph in enumerate(document.paragraphs):
            if "[python_fig3]" in paragraph.text:
                linje_til_fig3 = paragraph_index
                break
        document.paragraphs[linje_til_fig3].clear()
        run3 = document.paragraphs[linje_til_fig3].add_run()
        bilde3 = run3.add_picture("src/data/TRT-figurer//fig3.png")
        bilde3.width = Cm(16)
        bilde3.height = Cm(9)

        for paragraph_index, paragraph in enumerate(document.paragraphs):
            if "[python_fig5]" in paragraph.text:
                linje_til_fig5 = paragraph_index
                break
        document.paragraphs[linje_til_fig5].clear()
        run5 = document.paragraphs[linje_til_fig5].add_run()
        bilde5 = run5.add_picture("src/data/TRT-figurer//fig5.png")
        bilde5.width = Cm(16)
        bilde5.height = Cm(9)

        for paragraph_index, paragraph in enumerate(document.paragraphs):
            if "[python_fig6]" in paragraph.text:
                linje_til_fig6 = paragraph_index
                break
        document.paragraphs[linje_til_fig6].clear()
        run6 = document.paragraphs[linje_til_fig6].add_run()
        bilde6 = run6.add_picture("src/data/TRT-figurer//fig6.png")
        bilde6.width = Cm(16)
        bilde6.height = Cm(9)

        st.markdown("---")
        bio = io.BytesIO()
        document.save(bio)
        if document:
            st.download_button(
                label="Last ned rapport 📝",
                data=bio.getvalue(),
                file_name="TRT-rapport med figurer.docx",
                mime="docx")

TRT_beregning().kjor_hele()