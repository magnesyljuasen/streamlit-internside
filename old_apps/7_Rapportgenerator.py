import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
import io
import datetime

st.set_page_config(page_title="Rapportgenerator", page_icon="📘")

with open("styles/main.css") as f:
    st.markdown("<style>{}</style>".format(f.read()), unsafe_allow_html=True) 
    
st.title("Rapportgenerator")
st.button("Refresh")

document = Document("src/data/docx/Rapportmal.docx")
styles = document.styles
style = styles.add_style('Citation', WD_STYLE_TYPE.PARAGRAPH)

with st.form("Inndata"):
    c1, c2 = st.columns(2)
    with c1:
        TITTEL = st.text_input("Tittel", value = "Termisk responstest")
        OPPDRAGSNAVN = st.text_input("Oppdragsnavn", value = "Termisk responstest - Høgda")
        OPPDRAGSGIVER = st.text_input("Oppdragsgiver", value = "Båsum Boring")
        OPPDRAGSNUMMER = st.text_input("Oppdragsnummer", value = "630962-01")
        FORFATTER = st.text_input("Utarbeidet av", value = "Magne Syljuåsen")
        OPPDRAGSLEDER = st.text_input("Oppdragsleder", value = "Henrik Holmberg")
        KVALITETSSIKRER = st.text_input("Kvalitetssikrer", value = "Sofie Hartvigsen")
    with c2:
        KONTOR = st.text_input("Kontor", value = "Trondheim")
    st.form_submit_button("Gi input")

document.paragraphs[1].text = f"Oppdragsgiver:      {OPPDRAGSGIVER}"
document.paragraphs[2].text = f"Oppdragsgiver:      {OPPDRAGSGIVER}"
document.paragraphs[3].text = f"Tittel på rapport:      {OPPDRAGSNAVN}"
document.paragraphs[4].text = f"Tittel på rapport:      {OPPDRAGSNAVN}"
st.write("---")
"""
st.write(document.paragraphs[2].text)
st.write(document.paragraphs[3].text)
st.write(document.paragraphs[4].text)
st.write(document.paragraphs[5].text)
st.write(document.paragraphs[6].text)
st.write(document.paragraphs[7].text)
st.write(document.paragraphs[8].text)
st.write(document.paragraphs[9].text)
st.write(document.paragraphs[10].text)
st.write(document.paragraphs[11].text)
st.write(document.paragraphs[12].text)
st.write(document.paragraphs[13].text)
st.write(document.paragraphs[14].text)
st.write(document.paragraphs[15].text)
st.write(document.paragraphs[16].text)
st.write(document.paragraphs[17].text)
st.write(document.paragraphs[18].text)
st.write(document.paragraphs[19].text)
st.write(document.paragraphs[20].text)
st.write(document.paragraphs[21].text)
"""

#--


"""
st.subheader("Eksempel: Parametere")
with st.form("Input"):
    forfatter = st.text_input("Forfatter", value="Ola Nordmann")
    oppdragsleder = st.text_input("Oppdragsleder", value="Kari Nordmann")
    oppdragsgiver = st.text_input("Oppdragsgiver", value = "Firma AS")
    oppdragsnummer = st.text_input("Oppdragsnummer", value = "635960-01")
    sted = st.text_input("Sted", value = "Trondheim")
    #--
    depth_to_bedrock = st.number_input("Dybde til fjell [m]", value=5, step=1)
    loose_material = st.selectbox("Hva slags løsmasser?", options=["hav- og fjordavsetning", "elveavsetning", "breelvavsetning", "morene"])
    st.form_submit_button("Gi input")
#--
if depth_to_bedrock > 15:
    setning_dybde_til_fjell = "Siden dybde til fjell > 15 m, må det gjøres supplerende grunnundersøkelser. "
else:
    setning_dybde_til_fjell = ""
#--
document = Document("src/data/docx/Termisk responstest.docx")
styles = document.styles
style = styles.add_style('Citation', WD_STYLE_TYPE.PARAGRAPH)
document.paragraphs[0].text = f"Oppdragsgiver: {oppdragsgiver}"
document.paragraphs[1].text = f"Oppdragsnavn: Geoteknisk rapport - {sted}"
document.paragraphs[2].text = f"Oppdragsnummer: {oppdragsnummer} - {sted}"
document.paragraphs[3].text = f"Utarbeidet av: {forfatter}"
document.paragraphs[4].text = f"Oppdragsleder: {oppdragsleder}"
document.paragraphs[5].text = f"Dato: {datetime.date.today()}"
document.paragraphs[6].text = f"Tilgjenglighet: Åpen"

document.paragraphs[7].text = f"Geoteknisk notat - {sted}"

document.add_heading("Innledning", 1)
document.add_paragraph(report_text_1)
"""
#--
st.markdown("---")
bio = io.BytesIO()
document.save(bio)
if document:
    st.download_button(
        label="Last ned rapport!",
        data=bio.getvalue(),
        file_name="Rapport.docx",
        mime="docx")
